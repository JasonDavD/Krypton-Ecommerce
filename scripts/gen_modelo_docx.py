# -*- coding: utf-8 -*-
"""Genera Krypton_Modelo_Datos.docx (documento didactico del modelo de datos)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ACCENT = RGBColor(0x1F, 0x4E, 0x79)   # azul Krypton
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# --- estilos base ---
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = ACCENT
    return h


def para(text, italic=False, bold=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    r.font.size = Pt(size)
    return p


def mono(text, size=8):
    p = doc.add_paragraph()
    for line in text.split("\n"):
        run = p.add_run(line + "\n")
        run.font.name = "Consolas"
        run.font.size = Pt(size)
    return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].paragraphs[0].add_run(h).bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


# ====================== PORTADA ======================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("KRYPTON E-COMMERCE")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = ACCENT

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Modelo de Datos — Documento Didáctico")
r.font.size = Pt(16)
r.font.color.rgb = GREY

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(
    "E-commerce B2C de artefactos tecnológicos  ·  Spring Boot + PostgreSQL\n"
    "CIBERTEC — EFSRT VI"
).font.size = Pt(11)

doc.add_paragraph()
para(
    "Propósito: explicar el modelo de datos, las relaciones entre tablas y las "
    "decisiones de diseño, de forma que todo el equipo (codificación y "
    "documentación) entienda el qué y, sobre todo, el porqué.",
    italic=True,
)
doc.add_page_break()

# ====================== 1. VISIÓN GENERAL ======================
heading("1. Visión general", 1)
para("El modelo tiene 8 tablas. Vista general de las relaciones:")
mono(
    "CATEGORY ─1:N─> PRODUCT ─1:N─> STOCK_MOVEMENT\n"
    "                   |\n"
    "            +──────+──────+\n"
    "          1:N            1:N\n"
    "            |              |\n"
    "        CART_ITEM      ORDER_ITEM\n"
    "            |              |\n"
    "          N:1            N:1\n"
    "            |              |\n"
    "USER ─1:1─> CART       ORDER\n"
    "  |                       ^\n"
    "  +──────────1:N──────────+\n",
    size=9,
)
table(
    ["#", "Tabla", "Para qué existe"],
    [
        ["1", "users", "Clientes y administradores de la plataforma"],
        ["2", "categories", "Agrupa los productos (reportes por categoría)"],
        ["3", "products", "El catálogo: lo que se vende"],
        ["4", "cart", "El carrito activo de cada usuario registrado"],
        ["5", "cart_item", "Cada producto agregado al carrito"],
        ["6", "orders", "Las compras confirmadas"],
        ["7", "order_items", "El detalle (líneas) de cada compra"],
        ["8", "stock_movement", "Historial de entradas/salidas (kardex)"],
    ],
    widths=[0.4, 1.4, 4.2],
)

# ====================== 2. TABLAS EN DETALLE ======================
heading("2. Las tablas en detalle", 1)

tablas = {
    "users": [
        ["id", "bigint PK", "identificador técnico"],
        ["name", "varchar", "nombre del usuario"],
        ["email", "varchar UNIQUE", "login, no se repite"],
        ["password", "varchar", "hasheada (BCrypt), nunca en texto plano"],
        ["role", "varchar", "CLIENTE o ADMIN"],
        ["created_at", "timestamp", "fecha de registro"],
    ],
    "categories": [
        ["id", "bigint PK", ""],
        ["name", "varchar UNIQUE", "ej: Laptops, Celulares"],
        ["description", "varchar", ""],
    ],
    "products": [
        ["id", "bigint PK", "clave técnica (surrogate)"],
        ["sku", "varchar UNIQUE", "clave de negocio (LAP-DELL-15-001)"],
        ["name", "varchar", ""],
        ["description", "varchar", ""],
        ["price", "decimal", "precio actual"],
        ["stock", "int", "VALOR CACHEADO del stock (ver §4)"],
        ["image_url", "varchar", ""],
        ["active", "boolean", "si se muestra en el catálogo"],
        ["category_id", "bigint FK", "-> categories"],
    ],
    "cart": [
        ["id", "bigint PK", ""],
        ["user_id", "bigint FK UNIQUE", "1 carrito activo por usuario"],
        ["created_at / updated_at", "timestamp", ""],
    ],
    "cart_item": [
        ["id", "bigint PK", ""],
        ["cart_id", "bigint FK", "-> cart"],
        ["product_id", "bigint FK", "-> products"],
        ["quantity", "int", "no guarda precio: usa el precio VIVO"],
    ],
    "orders": [
        ["id", "bigint PK", ""],
        ["user_id", "bigint FK", "quién compró"],
        ["order_date", "timestamp", ""],
        ["status", "varchar", "PENDIENTE / CONFIRMADA / CANCELADA"],
        ["total", "decimal", "total de la compra"],
    ],
    "order_items": [
        ["id", "bigint PK", ""],
        ["order_id", "bigint FK", "-> orders"],
        ["product_id", "bigint FK", "-> products"],
        ["quantity", "int", ""],
        ["unit_price", "decimal", "SNAPSHOT: precio congelado al comprar"],
    ],
    "stock_movement (el kardex)": [
        ["id", "bigint PK", ""],
        ["product_id", "bigint FK", "-> products"],
        ["type", "varchar", "ENTRADA o SALIDA"],
        ["quantity", "int", "cantidad del movimiento"],
        ["reason", "varchar", "compra, venta, ajuste, devolución"],
        ["reference", "varchar", "referencia al origen (ej: order_id)"],
        ["created_at", "timestamp", ""],
        ["created_by", "bigint FK", "-> users (el admin)"],
    ],
}
for name, rows in tablas.items():
    heading(name, 2)
    table(["Campo", "Tipo", "Notas"], rows, widths=[1.7, 1.6, 3.0])

# ====================== 3. RELACIONES ======================
heading("3. Las relaciones, una por una", 1)
for txt in [
    "Category 1:N Product — una categoría agrupa muchos productos; cada producto "
    "pertenece a una sola categoría. Esto habilita los reportes por categoría.",
    "User 1:N Order — un cliente hace muchas órdenes; cada orden es de un cliente.",
    "User 1:1 Cart — cada usuario tiene un único carrito activo (user_id UNIQUE).",
    "Cart 1:N CartItem — un carrito tiene varias líneas, una por producto.",
    "Order 1:N OrderItem — una orden tiene una o más líneas de detalle.",
    "Product 1:N CartItem / OrderItem / StockMovement — un producto aparece en "
    "muchos carritos, muchas órdenes y tiene muchos movimientos de stock.",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(txt)

heading("¿Por qué existen cart_item y order_item?", 2)
para(
    "Entre Order y Product (y entre Cart y Product) hay una relación "
    "muchos-a-muchos: una orden tiene muchos productos y un producto está en "
    "muchas órdenes. En una base relacional un M:N no se puede representar "
    "directo: se resuelve con una tabla intermedia (entidad asociativa). Esa "
    "tabla es order_item (y cart_item). Además guardan datos propios de la "
    "relación (quantity, unit_price), así que son obligatorias."
)

# ====================== 4. CONCEPTOS CLAVE ======================
heading("4. Conceptos clave (la parte importante)", 1)

heading("4.1 Clave técnica vs. clave de negocio (id vs sku)", 2)
para(
    "products tiene dos identificadores. id es la surrogate key: autoincremental, "
    "interna, rápida; es la que se usa en todas las FKs. sku es la natural key "
    "(clave de negocio): única, legible por humanos (LAP-DELL-15-001), la que ve "
    "y busca el admin. Regla: nunca uses el SKU como FK; para conectar tablas, "
    "siempre el id."
)

heading("4.2 El snapshot de precio (precio vivo vs. congelado)", 2)
para(
    "cart_item usa precio VIVO: no guarda precio, el carrito siempre muestra el "
    "precio ACTUAL del producto. order_item usa precio CONGELADO: en el checkout "
    "el precio del momento se copia a unit_price y queda como una foto histórica."
)
para(
    "¿Por qué? El precio cambia con el tiempo. Si la orden solo referenciara al "
    "producto, al subir el precio mañana todas las órdenes viejas mostrarían el "
    "precio nuevo: un error contable grave. El checkout es la frontera entre el "
    "precio vivo y el precio congelado.",
    italic=True,
)

heading("4.3 El valor cacheado de products.stock (denormalización)", 2)
para("Con el kardex (stock_movement), el stock real se puede calcular:")
mono("stock real = SUM(ENTRADAS) − SUM(SALIDAS)\n"
     "ej: entró 100, salió 3, salió 2  ->  stock real = 95", size=10)
para(
    "Entonces, en teoría, no haría falta la columna stock: el dato verdadero vive "
    "en el historial. stock_movement es la fuente de la verdad. ¿Por qué guardamos "
    "igual el número en products.stock? Por PERFORMANCE: con miles de productos y "
    "cientos de movimientos cada uno, calcular ese SUM en cada carga del catálogo "
    "sería lentísimo. Guardamos el resultado ya calculado y leerlo es instantáneo. "
    "Eso es el valor cacheado (técnicamente, denormalización: guardar a propósito "
    "un dato redundante derivable, para ganar velocidad de lectura)."
)
ban = doc.add_paragraph()
ban.add_run("Analogía — tu cuenta bancaria. ").bold = True
ban.add_run(
    "El banco podría calcular tu saldo sumando cada movimiento desde que abriste "
    "la cuenta. No lo hace: guarda el saldo actual (valor cacheado) y lo actualiza "
    "con cada transacción. El historial es la verdad; el saldo es la lectura "
    "rápida. En Krypton: stock_movement = historial, products.stock = saldo."
).italic = True

heading("4.4 El checkout es UNA transacción atómica", 2)
para(
    "Cuando un cliente confirma la compra, dentro de una sola transacción pasa "
    "todo esto — o pasa todo, o no pasa nada (rollback):"
)
for i, step in enumerate([
    "Se crea la order y sus order_items (copiando el precio -> snapshot).",
    "Por cada ítem se genera una SALIDA en stock_movement.",
    "Se descuenta products.stock (el valor cacheado).",
    "Se vacía el cart.",
], 1):
    p = doc.add_paragraph(style="List Number")
    p.add_run(step)
para(
    "Si cualquier paso falla (ej: no hay stock suficiente), se revierte TODO. Así "
    "stock_movement y products.stock siempre cuadran, y nunca se vende algo sin "
    "descontarlo. Esto cumple el requisito de seguridad transaccional de la rúbrica.",
    bold=True,
)

# ====================== 5. GLOSARIO ======================
heading("5. Glosario rápido", 1)
table(
    ["Término", "Qué es"],
    [
        ["PK (Primary Key)", "identificador único de la fila"],
        ["FK (Foreign Key)", "columna que apunta a la PK de otra tabla"],
        ["UNIQUE", "valor que no se puede repetir en la tabla"],
        ["Surrogate key", "PK artificial, autoincremental (id)"],
        ["Natural / business key", "identificador con significado de negocio (sku)"],
        ["Entidad asociativa", "tabla intermedia que resuelve un M:N (order_item)"],
        ["Snapshot", "copia congelada de un valor en un momento (unit_price)"],
        ["Denormalización", "guardar un dato redundante a propósito, por performance"],
        ["Kardex", "registro histórico de entradas/salidas de inventario"],
        ["Transacción atómica", "operaciones que pasan todas o ninguna"],
    ],
    widths=[1.9, 4.4],
)

out = r"F:\CIBERTEC\VI\EFSRT V\Proyecto\Krypton-Ecommerce\Krypton_Modelo_Datos.docx"
doc.save(out)
print("OK ->", out)
